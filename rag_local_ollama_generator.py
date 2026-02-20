from pathlib import Path
from uuid import uuid4

import chromadb
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from pypdf import PdfReader

DATA_DIR = Path("data")
DB_DIR = "chroma_db"
COLLECTION_NAME = "local_rag_demo"
EMBEDDING_MODEL = "nomic-embed-text"


def load_text_docs() -> list[Document]:
    print("[load_text_docs] Start")
    docs: list[Document] = []
    if not DATA_DIR.exists():
        print("[load_text_docs] data/ folder not found, skipping txt files")
        return docs

    for p in DATA_DIR.glob("*.txt"):
        print(f"[load_text_docs] Reading: {p}")
        text = p.read_text(encoding="utf-8").strip()
        if text:
            docs.append(Document(page_content=text, metadata={"source": str(p)}))
            print(f"[load_text_docs] Added txt doc: {p}")
    print(f"[load_text_docs] Completed. Total txt docs: {len(docs)}")
    return docs


def load_docx_docs() -> list[Document]:
    print("[load_docx_docs] Start")
    docs: list[Document] = []
    docx_paths = list(DATA_DIR.glob("*.docx")) if DATA_DIR.exists() else []
    print(f"[load_docx_docs] Found docx files: {len(docx_paths)}")

    # Keep unique files in case paths overlap.
    unique_paths = {str(p.resolve()): p for p in docx_paths}.values()

    for p in unique_paths:
        print(f"[load_docx_docs] Reading: {p}")
        doc = DocxDocument(str(p))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
        if text:
            docs.append(Document(page_content=text, metadata={"source": str(p)}))
            print(f"[load_docx_docs] Added docx doc: {p}")
    print(f"[load_docx_docs] Completed. Total docx docs: {len(docs)}")
    return docs


def load_pdf_docs() -> list[Document]:
    print("[load_pdf_docs] Start")
    docs: list[Document] = []
    pdf_paths = list(DATA_DIR.glob("*.pdf")) if DATA_DIR.exists() else []
    print(f"[load_pdf_docs] Found pdf files: {len(pdf_paths)}")

    for pdf_path in pdf_paths:
        print(f"[load_pdf_docs] Reading: {pdf_path}")
        reader = PdfReader(str(pdf_path))
        print(f"[load_pdf_docs] Pages detected in {pdf_path.name}: {len(reader.pages)}")
        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                print(f"[load_pdf_docs] Skipping empty page: {pdf_path.name} page {idx}")
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": str(pdf_path), "page": idx},
                )
            )
            print(f"[load_pdf_docs] Added page: {pdf_path.name} page {idx}")
    print(f"[load_pdf_docs] Completed. Total pdf docs: {len(docs)}")
    return docs


def load_docs() -> list[Document]:
    print("[load_docs] Start")
    if not DATA_DIR.exists():
        raise RuntimeError("data/ folder not found.")

    docs: list[Document] = []
    docs.extend(load_text_docs())
    docs.extend(load_docx_docs())
    docs.extend(load_pdf_docs())
    if not docs:
        raise RuntimeError("No input documents found in data/ (.txt/.docx/.pdf).")
    print(f"[load_docs] Completed. Total docs loaded: {len(docs)}")
    return docs


def reset_collection() -> None:
    print("[reset_collection] Start")
    client = chromadb.PersistentClient(path=DB_DIR)

    try:
        existing = client.get_collection(COLLECTION_NAME)
        existing_count = existing.count()
        print(f"[reset_collection] Existing indexes before reset: {existing_count}")
        client.delete_collection(COLLECTION_NAME)
        print("[reset_collection] Existing collection deleted")
    except Exception:
        print("[reset_collection] No existing collection found")

    client.create_collection(COLLECTION_NAME)
    empty_count = client.get_collection(COLLECTION_NAME).count()
    print(f"[reset_collection] Collection recreated. Indexes now: {empty_count}")
    print("[reset_collection] There are no index present in ChromaDB.")


def build_and_persist_index() -> None:
    print("[build_and_persist_index] Start")
    reset_collection()
    docs = load_docs()
    print("[build_and_persist_index] Splitting documents into chunks")
    splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50)
    chunks = splitter.split_documents(docs)
    print(f"[build_and_persist_index] Chunking complete. Total chunks: {len(chunks)}")

    print(f"[build_and_persist_index] Initializing embedding model: {EMBEDDING_MODEL}")
    embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL)
    print("[build_and_persist_index] Initializing Chroma collection")
    db = Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=DB_DIR,
        embedding_function=embeddings,
    )
    print("[build_and_persist_index] Chroma initialized")

    print("[build_and_persist_index] Indexing chunks one by one")
    for idx, chunk in enumerate(chunks, start=1):
        db.add_documents([chunk], ids=[str(uuid4())])
        print(f"[index_progress] Indexed chunk {idx}/{len(chunks)}")

    print(f"Indexed {len(docs)} docs into {len(chunks)} chunks.")

    # Reload from persisted vector DB and print one stored chunk.
    print("[build_and_persist_index] Reloading Chroma collection for verification")
    db_verify = Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=DB_DIR,
        embedding_function=embeddings,
    )
    total_indexes = db_verify._collection.count()
    print(f"Total indexes in vector DB: {total_indexes}")

    print("[build_and_persist_index] Fetching one sample chunk")
    sample = db_verify.get(limit=1, include=["documents", "metadatas"])
    sample_docs = sample.get("documents", [])
    sample_meta = sample.get("metadatas", [])

    if sample_docs:
        print("\nOne chunk from vector DB:")
        print(sample_docs[0][:500])
        if sample_meta:
            print("Metadata:", sample_meta[0])
    else:
        print("No chunk found in vector DB.")
    print("[build_and_persist_index] Completed")


if __name__ == "__main__":
    build_and_persist_index()
